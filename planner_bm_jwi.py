import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# --- 1. CONFIGURATION & CORE SETTINGS ---
st.set_page_config(page_title="Penjana RMH IRK Brunei (2047)", layout="wide")
st.title("🕌 RANCANGAN MENGAJAR HARIAN PUI/IRK 2047")

user_api_key = st.text_input("🔑 MASUKKAN KUNCI API GEMINI ANDA:", type="password")

def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return None
    return "models/gemini-1.5-flash"

selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)


# --- 2. DUAL-LANGUAGE AI ENGINE ROUTINES ---
def generate_irk_plan(topic, extra_context, api_key, model_name, language_mode):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    if language_mode == 'JAWI':
        prompt = f"""
        Topik/Tajuk: {topic}. 
        Sukatan Pelajaran: IRK Brunei Syllabus Code 2047 (Cambridge GCE O Level).
        Konteks Tambahan: {extra_context}.
        
        Sila hasilkan Rancangan Mengajar Harian (RMH) yang lengkap dan profesional dalam TULISAN JAWI sepenuhnya.
        
        PERATURAN FORMAT KRITIKAL:
        1. JANGAN gunakan rumi tulisan biasa, abjad inggeris (seperti perkataan 'minit') atau tanda asteris (**) sama sekali dalam kandungan. 
        2. Sila gunakan ejaan Jawi 'مينيت' untuk menyatakan tempoh masa. Contoh: (١٠ مينيت) atau (٥ مينيت).
        3. Gunakan nombor Jawi/Arab sah seperti (١, ٢, ٣, ٤, ٥...) untuk semua jenis senarai. JANGAN GUNA (1, 2, 3...).
        4. Setiap tajuk bahagian WAJIB dimulakan dengan perkataan "بهاڬين: " diikuti nama bahagian tersebut dalam Jawi.
        5. JANGAN guna perkataan MURID, digantikan dengan perkataan PELAJAR menyeluruh di dalam teks.
        
        Strukturkan output tepat mengikut penanda berikut:
        
        بهاڬين: توڤيک دان كود سوكتن
        {topic} (IRK BRUNEI 2047)
        
        بهاڬين: اوبجيكتيف ڤمبلاجرن
        [Tuliskan 4 mata menggunakan penomboran ١ hingga ٤]
        
        بهاڬين: حاصيل ڤمبلاجرن
        [Tuliskan 4 mata menggunakan penomboran ١ hingga ٤]
        
        بهاڬين: دليل نقلي دان روجوعن كتاب
        [Tuliskan ٢ dalil/rujukan menggunakan penomboran ١ dan ٢]
        
        بهاڬين: كريتيريا كجايان
        [Tuliskan 4 mata menggunakan penomboran ١ hingga ٤]
        
        بهاڬين: كات چونچي شريعة دان اصطلاح
        [Sediakan kata kunci dipisahkan dengan tanda koma Arab '،' tanpa sebarang nombor siri di hadapannya]
        
        بهاڬين: سوءالن كمهيرن برفكير ارس تيڠڬي
        [Tuliskan 4 soalan KBAT menggunakan penomboran ١ hingga ٤]
        
        بهاڬين: ستراتيڬي ڤڠاجرن تربيزا
        ١. ڤلاجر تاهاڤ تيڠڬي (HA - هيجاو): [Sediakan 1 aktiviti spesifik]
        ٢. ڤلاجر تاهاڤ سدرهان (MA - كونيڠ): [Sediakan 1 aktiviti spesifik]
        ٣. ڤلاجر تاهاڤ رنده (LA - ميره): [Sediakan 1 aktiviti spesifik]
        
        بهاڬين: اكتيۏيتي ڤمبلاجرن اوتام
        [Tuliskan langkah-langkah menggunakan penomboran ١، ٢، ٣...]
        
        بهاڬين: ڤنيلاين دان توڬاسن رومه
        [Tuliskan tugasan menggunakan penomboran ١ dan ٢]
        """
    else:
        # Standard Rumi Prompt with matching tables and parameters
        prompt = f"""
        Topik/Tajuk: {topic}. Sukatan Pelajaran: IRK Brunei 2047. Konteks Tambahan: {extra_context}.
        Sila hasilkan RMH lengkap dalam Bahasa Melayu Rumi standard. Pastikan menggunakan istilah PELAJAR (bukan MURID).
        JANGAN gunakan penomboran siri (1, 2, 3) di dalam bahagian KATA KUNCI SYARIAT / ISTILAH. Sila pisahkan kata kunci tersebut menggunakan tanda koma (,).
        Gunakan penanda siri "SECTION: " diikuti nama bahagian dalam huruf besar. Gunakan angka biasa (1, 2, 3).
        
        SECTION: TOPIK DAN KOD SUKATAN
        SECTION: OBJEKTIF PEMBELAJARAN
        SECTION: HASIL PEMBELAJARAN
        SECTION: DALIL NAQLI DAN RUJUKAN KITAB
        SECTION: KRITERIA KEJAYAAN
        SECTION: KATA KUNCI SYARIAT / ISTILAH
        SECTION: SOALAN KEMAHIRAN BERFIKIR ARAS TINGGI (KBAT)
        SECTION: STRATEGI PENGAJARAN TERBEZA
        1. Pelajar Tahap Tinggi (HA - Hijau): [Sediakan 1 aktiviti spesifik]
        2. Pelajar Tahap Sederhana (MA - Kuning): [Sediakan 1 aktiviti spesifik]
        3. Pelajar Tahap Rendah (LA - Merah): [Sediakan 1 aktiviti spesifik]
        
        SECTION: AKTIVITI PEMBELAJARAN UTAMA
        SECTION: PENILAIAN & TUGASAN RUMAH
        """
        
    try:
        response = model.generate_content(prompt)
        return response.text.replace("**", "")
    except Exception as e:
        return f"RALAT: {str(e)}"


# --- 3. EXPORT GENERATION UTILITIES ---
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_word_export(topic, text, is_jawi=False):
    doc = Document()
    
    # Page Margins Setup
    section_geo = doc.sections[0]
    section_geo.page_width = Inches(8.5)
    section_geo.page_height = Inches(11.5)
    for m in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
        setattr(section_geo, m, Inches(0.4))
    
    # Clean Running Header (Numeric values centered)
    header_p = section_geo.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run()
    add_page_number(header_run)
    header_run.font.name = 'Arial'
    header_run.font.size = Pt(11)

    # Styles Config
    style = doc.styles['Normal']
    style.font.name = 'Traditional Arabic' if is_jawi else 'Arial'
    style.font.size = Pt(16) if is_jawi else Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(10)

    # Title Banner
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_jawi else WD_ALIGN_PARAGRAPH.LEFT
    title_text = f"رانچڠن مڠاجر هارين: {topic}" if is_jawi else f"RANCANGAN MENGAJAR HARIAN: {topic.upper()}"
    title_p.add_run(title_text).bold = True

    # Administrative Table Grid Setup
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    
    if is_jawi:
        jawi_labels = [
            ["تاريق:", "هاري:"],
            ["تمڤت / بيليک:", "ميڠڬو نو:"],
            ["بيلڠن ڤلاجر:", "تيمڤوه (مينيت):"]
        ]
        for r in range(3):
            cell_a = admin_table.cell(r, 1)
            p_a = cell_a.paragraphs[0]
            p_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_a.add_run(jawi_labels[r][0]).bold = True
            
            cell_b = admin_table.cell(r, 3)
            p_b = cell_b.paragraphs[0]
            p_b.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_b.add_run(jawi_labels[r][1]).bold = True
    else:
        rumi_labels = [
            ["TARIKH:", "HARI:"],
            ["TEMPAT / BILIK:", "MINGGU NO:"],
            ["BILANGAN PELAJAR:", "TEMPOH (MINIT):"]
        ]
        for r in range(3):
            cell_a = admin_table.cell(r, 0)
            p_a = cell_a.paragraphs[0]
            p_a.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_a.add_run(rumi_labels[r][0]).bold = True
            
            cell_b = admin_table.cell(r, 2)
            p_b = cell_b.paragraphs[0]
            p_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_b.add_run(rumi_labels[r][1]).bold = True

    doc.add_paragraph()

    # Section Splitting Criteria Selection
    delimiter = 'بهاڬين:' if is_jawi else 'SECTION:'
    sections = text.split(delimiter)
    
    for section in sections:
        if not section.strip():
            continue
            
        lines = section.strip().split('\n')
        title = lines[0].strip()
        body_content = "\n".join(lines[1:]).strip()

        # Section Heading
        p_sec = doc.add_paragraph()
        p_sec.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_jawi else WD_ALIGN_PARAGRAPH.LEFT
        p_sec.add_run(title).bold = True

        # Custom Table Renderer: Universal Grid Maker for both Rumi/Jawi Keywords
        if ("اصطلاح" in title or "KEYWORD" in title or "KATA KUNCI" in title):
            delimit_char = '،' if is_jawi else ','
            keywords = [k.strip() for k in body_content.split(delimit_char) if k.strip()]
            
            # Determine dynamic row counts for long keyword lists (e.g. 12 items -> 4 rows x 3 columns)
            num_keywords = len(keywords)
            num_cols = 3
            num_rows = max(2, (num_keywords + num_cols - 1) // num_cols)
            
            while len(keywords) < (num_rows * num_cols):
                keywords.append("-")
                
            kw_table = doc.add_table(rows=num_rows, cols=num_cols)
            kw_table.style = 'Table Grid'
            
            kw_idx = 0
            for r_kw in range(num_rows):
                for c_kw in range(num_cols):
                    kw_cell = kw_table.cell(r_kw, c_kw)
                    kw_p = kw_cell.paragraphs[0]
                    kw_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    kw_p.add_run(keywords[kw_idx])
                    kw_idx += 1
            doc.add_paragraph()
            continue

        # Standard Text Boxes for other sections
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell_p = table.cell(0, 0).paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_jawi else WD_ALIGN_PARAGRAPH.LEFT
        cell_p.add_run(body_content if body_content else "...")
        doc.add_paragraph()

    # Dynamic Evaluation Footer Block
    doc.add_page_break()
    p_hod = doc.add_paragraph()
    p_hod.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_jawi else WD_ALIGN_PARAGRAPH.LEFT
    p_hod.add_run("ڤڠسهن دان اولاسن ڤڠتوا / كتوا جابتن" if is_jawi else "PENGESAHAN DAN ULASAN PENGETUA / KETUA JABATAN").bold = True
    
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    
    if is_jawi:
        hod_table.cell(0, 0).paragraphs[0].add_run(" :اولاسن").bold = True
        hod_table.cell(0, 1).paragraphs[0].add_run(" :تنداتاڠن دان چوڤ جابتن").bold = True
        hod_table.rows[1].height = Pt(55)
        hod_table.cell(2, 0).paragraphs[0].add_run(" :تاريق سيمقن").bold = True
        hod_table.cell(2, 1).paragraphs[0].add_run(" :نام ڤڽيمق").bold = True
        for row in hod_table.rows:
            for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        hod_table.cell(0, 0).paragraphs[0].add_run("ULASAN / REMARK:").bold = True
        hod_table.cell(0, 1).paragraphs[0].add_run("TANDATANGAN & COP JABATAN:").bold = True
        hod_table.rows[1].height = Pt(50)
        hod_table.cell(2, 0).paragraphs[0].add_run("TARIKH SEMAKAN:").bold = True
        hod_table.cell(2, 1).paragraphs[0].add_run("NAMA PENYEMAK:").bold = True

    # Layout Spacing Adjustments
    for row in admin_table.rows:
        for cell in row.cells: cell.paragraphs[0].paragraph_format.line_spacing = 1.0
    for row in hod_table.rows:
        for cell in row.cells: cell.paragraphs[0].paragraph_format.line_spacing = 1.0

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 4. STREAMLIT INTERFACE ---
st.write("---")
u_topic = st.text_input("TOPIK / TAJUK PELAJARAN IRK:", value="Rukun Sembahyang Lima Waktu")
u_extra = st.text_area("KONTEKS TAMBAHAN / NOTA KHUSUS (PILIHAN):")

if st.button("🚀 KLIK JANA DUA VERSI Rancangan Mengajar (RUMI & JAWI)", type="primary"):
    if not user_api_key:
        st.error("❌ SILA MASUKKAN KUNCI API GEMINI ANDA DI BAHAGIAN ATAS.")
    elif not u_topic:
        st.error("❌ SILA ISI RUANGAN TOPIK.")
    else:
        with st.spinner("Menjana versi Rumi..."):
            st.session_state['irk_malay_out'] = generate_irk_plan(u_topic, u_extra, user_api_key, selected_model_name, "MALAY")
        with st.spinner("Menjana versi Jawi..."):
            st.session_state['irk_jawi_out'] = generate_irk_plan(u_topic, u_extra, user_api_key, selected_model_name, "JAWI")

if 'irk_malay_out' in st.session_state and 'irk_jawi_out' in st.session_state:
    st.divider()
    col1, col2 = st.columns(2)
    
    formatted_filename = u_topic.strip().replace(' ', '_')
    
    with col1:
        st.markdown("### 📝 Versi RMH BM")
        st.text_area("", st.session_state['irk_malay_out'], height=400)
        
        malay_doc = create_word_export(u_topic, st.session_state['irk_malay_out'], is_jawi=False)
        st.download_button(
            label="📥 DOWNLOAD WORD: Rancangan Mengajar Harian RUMI (.DOCX)",
            data=malay_doc, 
            file_name=f"RMH_IRK_RUMI_{formatted_filename}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    with col2:
        st.markdown("### 🕌 Versi RMH JAWI")
        st.markdown("")
        
        st.markdown(
            f'<textarea style="width:100%; height:400px; direction:rtl; text-align:right; font-family:\'Traditional Arabic\', sans-serif; font-size:18px;" readonly>{st.session_state["irk_jawi_out"]}</textarea>', 
            unsafe_allow_html=True
        )
        
        jawi_doc = create_word_export(u_topic, st.session_state['irk_jawi_out'], is_jawi=True)
        st.download_button(
            label="📥 DOWNLOAD WORD: Rancangan Mengajar Harian JAWI (.DOCX)",
            data=jawi_doc, 
            file_name=f"RMH_IRK_JAWI_{formatted_filename}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- UI FOOTNOTE BLOCKS (CENTERED) ---
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; color: #808495; font-size: 0.85rem;">
        Pencipta: Hjh Nurul Haziqah Hj Nordin | &copy; © 2026 BScHM Computer Science Strathclyde University
        Portal Penjana Rancangan Mengajar Harian ini untuk kegunaan Pegawai Pendidikan Brunei Darussalam sahaja.
    </div>
    """,
    unsafe_allow_html=True
)
